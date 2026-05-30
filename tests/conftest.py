from __future__ import annotations

import os
import sys
from pathlib import Path

import fitz
import pytest
from fastapi.testclient import TestClient
from docx import Document as DocxDocument
from openpyxl import Workbook
from sqlalchemy.orm import Session


PROJECT_ROOT = Path(__file__).resolve().parents[1]
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))


TEST_DB_PATH = Path("/private/tmp/oa_rag_test.sqlite3")
TEST_STORAGE_ROOT = Path("/private/tmp/oa_rag_test_data")

os.environ["DATABASE_URL"] = f"sqlite:///{TEST_DB_PATH}"
os.environ["STORAGE_ROOT"] = str(TEST_STORAGE_ROOT)
os.environ["RAW_DATA_DIR"] = str(TEST_STORAGE_ROOT / "raw")
os.environ["PROCESSED_DATA_DIR"] = str(TEST_STORAGE_ROOT / "processed")
os.environ["SAMPLE_DATA_DIR"] = str(TEST_STORAGE_ROOT / "samples")
os.environ["APP_API_KEY"] = "test-api-key"
os.environ["INTERNAL_TOKEN"] = "test-internal-token"
os.environ["ADMIN_MANUAL_USER_CREATION_ENABLED"] = "true"
os.environ["LEGACY_ADMIN_ROLE_SEARCH_BYPASS_ENABLED"] = "false"
os.environ["LEGACY_INTERNAL_WITHOUT_OWNER_ACCESS_ENABLED"] = "false"

from app.core.config import get_settings, reset_settings_cache
from app.db.base import Base
from app.db.runtime import get_engine, get_session_factory, reset_db_runtime
from app.integrations.vector_store import VectorStoreClient
from app.main import create_app
from app.retrieval.sparse_index import SparseIndexProvider
from app.services.background_jobs import BackgroundJobRunner


@pytest.fixture(autouse=True)
def reset_environment(monkeypatch):
    monkeypatch.setenv("DATABASE_URL", f"sqlite:///{TEST_DB_PATH}")
    monkeypatch.setenv("STORAGE_ROOT", str(TEST_STORAGE_ROOT))
    monkeypatch.setenv("RAW_DATA_DIR", str(TEST_STORAGE_ROOT / "raw"))
    monkeypatch.setenv("PROCESSED_DATA_DIR", str(TEST_STORAGE_ROOT / "processed"))
    monkeypatch.setenv("SAMPLE_DATA_DIR", str(TEST_STORAGE_ROOT / "samples"))
    monkeypatch.setenv("APP_API_KEY", "test-api-key")
    monkeypatch.setenv("INTERNAL_TOKEN", "test-internal-token")
    monkeypatch.setenv("ADMIN_MANUAL_USER_CREATION_ENABLED", "true")
    monkeypatch.setenv("LEGACY_ADMIN_ROLE_SEARCH_BYPASS_ENABLED", "false")
    monkeypatch.setenv("LEGACY_INTERNAL_WITHOUT_OWNER_ACCESS_ENABLED", "false")
    monkeypatch.setenv("RETRIEVAL_AUTHENTICATED_IDENTITY_REQUIRED", "false")
    monkeypatch.setenv("IDENTITY_SYNC_SCHEDULER_ENABLED", "false")
    monkeypatch.setenv("IDENTITY_SYNC_SCHEDULER_INTERVAL_SECONDS", "60")
    monkeypatch.setenv("IDENTITY_SYNC_SCHEDULER_BATCH_LIMIT", "20")
    monkeypatch.setenv("AUTH_SESSION_ENABLED", "false")
    monkeypatch.setenv("AUTH_SESSION_SECRET", "")
    monkeypatch.setenv("AUTH_SESSION_TTL_SECONDS", "3600")
    monkeypatch.setenv("AUTH_SESSION_ISSUER", "knowledge-base-core")
    monkeypatch.setenv("TRUSTED_IDENTITY_HEADER_ENABLED", "false")
    monkeypatch.setenv("TRUSTED_IDENTITY_EXTERNAL_SOURCE_HEADERS", "X-Identity-Source,X-Auth-Issuer")
    monkeypatch.setenv("TRUSTED_IDENTITY_EXTERNAL_ID_HEADERS", "X-Identity-External-Id,X-Auth-Subject")
    monkeypatch.setenv("TRUSTED_IDENTITY_JIT_ENABLED", "false")
    monkeypatch.setenv("EVALUATION_RETRIEVAL_USER_ID", "evaluation-service")
    monkeypatch.setenv("VECTOR_STORE_PROVIDER", "local")
    monkeypatch.setenv("EMBEDDING_PROVIDER", "local")
    monkeypatch.setenv("EMBEDDING_API_BASE", "")
    monkeypatch.setenv("EMBEDDING_API_KEY", "")
    monkeypatch.setenv("EMBEDDING_MODEL", "")
    monkeypatch.setenv("LLM_API_BASE", "")
    monkeypatch.setenv("LLM_API_KEY", "")
    monkeypatch.setenv("LLM_MODEL", "")
    monkeypatch.setenv("ZILLIZ_URI", "")
    monkeypatch.setenv("ZILLIZ_TOKEN", "")
    if TEST_DB_PATH.exists():
        TEST_DB_PATH.unlink()
    if TEST_STORAGE_ROOT.exists():
        for path in sorted(TEST_STORAGE_ROOT.rglob("*"), reverse=True):
            if path.is_file():
                path.unlink()
            elif path.is_dir():
                path.rmdir()
        TEST_STORAGE_ROOT.rmdir()

    reset_settings_cache()
    reset_db_runtime()
    VectorStoreClient._local_store.clear()
    SparseIndexProvider().build_index([])
    settings = get_settings()
    settings.ensure_storage_dirs()
    Base.metadata.create_all(bind=get_engine())
    yield
    BackgroundJobRunner.shutdown()
    VectorStoreClient._local_store.clear()
    SparseIndexProvider().build_index([])
    reset_db_runtime()
    reset_settings_cache()


@pytest.fixture
def client() -> TestClient:
    app = create_app()
    return TestClient(app)


@pytest.fixture
def db_session() -> Session:
    session = get_session_factory()()
    try:
        yield session
    finally:
        session.close()


@pytest.fixture
def sample_pdf_bytes() -> bytes:
    document = fitz.open()
    page = document.new_page()
    page.insert_text((72, 72), "OA rule update\nTransfer approval now requires two levels.")
    payload = document.tobytes()
    document.close()
    return payload


@pytest.fixture
def sample_docx_bytes(tmp_path) -> bytes:
    file_path = tmp_path / "sample.docx"
    document = DocxDocument()
    document.add_heading("调岗规则", level=1)
    document.add_paragraph("调岗申请需要增加二级审批。")
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "字段"
    table.cell(0, 1).text = "说明"
    table.cell(1, 0).text = "生效日期"
    table.cell(1, 1).text = "审批通过后的次月1日"
    document.save(file_path)
    return file_path.read_bytes()


@pytest.fixture
def sample_xlsx_bytes(tmp_path) -> bytes:
    file_path = tmp_path / "sample.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Config"
    worksheet.append(["key", "value"])
    worksheet.append(["approval_level", "2"])
    worksheet.append(["effective_rule", "next_month_first_day"])
    workbook.save(file_path)
    return file_path.read_bytes()


@pytest.fixture
def sample_xlsx_multiline_bytes(tmp_path) -> bytes:
    file_path = tmp_path / "sample_multiline.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "FAQ"
    worksheet.append(["一级", "二级", "三级"])
    worksheet.append(["基本咨询", "上传头像失败", "头像最多24张"])
    worksheet.append(["修改头像", "电脑端操作\n手机端操作", "如何更换头像"])
    worksheet.append(["删除头像", "联系人工处理", "删除头像和照片"])
    workbook.save(file_path)
    return file_path.read_bytes()


@pytest.fixture
def sample_xlsx_hierarchical_bytes(tmp_path) -> bytes:
    file_path = tmp_path / "sample_hierarchical.xlsx"
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "FAQ"
    worksheet.append(["一级", "二级", "三级", "四级", "相似问题1"])
    worksheet.append(["基本咨询", "上传头像", "照片上传失败", "请检查格式与大小", "头像上传不了"])
    worksheet.append([None, None, "如何修改头像", "请在个人资料中修改", "怎么换头像"])
    worksheet.append([None, "修改个人信息", "昵称如何修改", "请联系人工处理", "怎么改名字"])
    worksheet.append([None, None, "年龄如何修改", "请上传身份证明", "修改年龄"])
    workbook.save(file_path)
    return file_path.read_bytes()
