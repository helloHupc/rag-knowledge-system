from __future__ import annotations

import json
import os
import sys
from io import BytesIO
from pathlib import Path
from uuid import uuid4

from fastapi.testclient import TestClient


def build_docx_bytes() -> bytes:
    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading("审批流程说明", level=1)
    document.add_paragraph("OA 文档说明：审批流程需要两级审批。")
    document.add_paragraph("第一步由部门负责人审批，第二步由 HR 系统负责人审批。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def sync_smoke_identity(client: TestClient) -> bool:
    source_response = client.post(
        "/api/v1/admin/identity-sources",
        json={
            "source_id": "dify-smoke-directory",
            "source_type": "custom_api",
            "name": "Dify Smoke Directory",
            "field_mapping_json": {
                "user_id": "employee_no",
                "external_id": "directory_id",
                "display_name": "full_name",
                "email": "mail",
                "department_codes": "dept_codes",
                "status": "state",
            },
        },
    )
    if source_response.status_code not in {200, 409}:
        print(
            json.dumps(
                {"step": "create_identity_source", "status_code": source_response.status_code, "body": source_response.json()},
                ensure_ascii=False,
                indent=2,
            )
        )
        return False

    sync_response = client.post(
        "/api/v1/admin/identity-sources/dify-smoke-directory/sync",
        json={
            "records": [
                {
                    "employee_no": "dify-smoke",
                    "directory_id": "dify-smoke-user",
                    "full_name": "Dify Smoke User",
                    "mail": "dify-smoke@example.com",
                    "dept_codes": ["hr"],
                    "state": "active",
                },
                {
                    "employee_no": "dify-external",
                    "directory_id": "dify-external-service",
                    "full_name": "Dify External Retrieval Service",
                    "mail": "dify-external@example.com",
                    "dept_codes": ["hr"],
                    "state": "active",
                },
            ]
        },
    )
    if sync_response.status_code != 200:
        print(
            json.dumps(
                {"step": "sync_identity_source", "status_code": sync_response.status_code, "body": sync_response.json()},
                ensure_ascii=False,
                indent=2,
            )
        )
        return False
    return sync_response.json()["data"]["success_count"] == 2


def main() -> int:
    project_root = Path(__file__).resolve().parents[1]
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))

    from app.core.config import get_settings, reset_settings_cache

    reset_settings_cache()
    dify_app_key = os.environ.get("DIFY_APP_KEY", "").strip() or (get_settings().dify_app_key or "").strip()
    if not dify_app_key:
        print("DIFY_APP_KEY is required", file=sys.stderr)
        return 1

    smoke_id = uuid4().hex[:8]
    test_db_path = Path(f"/private/tmp/oa_rag_dify_smoke_{smoke_id}.sqlite3")
    test_storage_root = Path(f"/private/tmp/oa_rag_dify_smoke_data_{smoke_id}")

    os.environ["DATABASE_URL"] = f"sqlite:///{test_db_path}"
    os.environ["STORAGE_ROOT"] = str(test_storage_root)
    os.environ["RAW_DATA_DIR"] = str(test_storage_root / "raw")
    os.environ["PROCESSED_DATA_DIR"] = str(test_storage_root / "processed")
    os.environ["SAMPLE_DATA_DIR"] = str(test_storage_root / "samples")
    os.environ["INTERNAL_TOKEN"] = "smoke-internal-token"
    os.environ["VECTOR_STORE_PROVIDER"] = "local"
    os.environ["EMBEDDING_PROVIDER"] = "local"
    os.environ["EMBEDDING_API_BASE"] = ""
    os.environ["EMBEDDING_API_KEY"] = ""
    os.environ["EMBEDDING_MODEL"] = ""
    os.environ["LLM_API_BASE"] = ""
    os.environ["LLM_API_KEY"] = ""
    os.environ["LLM_MODEL"] = ""
    os.environ["ZILLIZ_URI"] = ""
    os.environ["ZILLIZ_TOKEN"] = ""

    from app.db.base import Base
    from app.db.runtime import get_engine, reset_db_runtime
    from app.main import create_app

    reset_settings_cache()
    reset_db_runtime()
    settings = get_settings()
    settings.ensure_storage_dirs()
    Base.metadata.create_all(bind=get_engine())

    app = create_app()
    client = TestClient(app)

    if not sync_smoke_identity(client):
        return 1

    upload_response = client.post(
        "/api/v1/documents/upload",
        data={
            "title": f"Dify Smoke {smoke_id}",
            "source_type": "rule_doc",
            "source_module": "oa",
            "version": smoke_id,
            "access_level": "internal",
            "owner_dept": "hr",
        },
        files={
            "file": (
                f"dify-smoke-{smoke_id}.docx",
                build_docx_bytes(),
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    if upload_response.status_code != 200:
        print(json.dumps({"status_code": upload_response.status_code, "body": upload_response.json()}, ensure_ascii=False, indent=2))
        return 1

    payload = {
        "query": "OA 文档里审批流程怎么说？",
        "response_mode": "qa",
        "top_k": 5,
        "filters": {"source_module": ["oa"]},
        "generation_options": {"temperature": 0.2, "max_tokens": 300},
    }
    response = client.post(
        "/api/v1/dify/knowledge",
        headers={"Authorization": f"Bearer {dify_app_key}"},
        json=payload,
    )
    external_response = client.post(
        "/api/v1/dify/retrieval",
        headers={"Authorization": f"Bearer {dify_app_key}"},
        json={
            "knowledge_id": "oa",
            "query": "两级审批",
            "retrieval_setting": {
                "top_k": 3,
                "score_threshold": 0.1,
            },
        },
    )
    result = {
        "knowledge": {"status_code": response.status_code, "body": response.json()},
        "external_retrieval": {
            "status_code": external_response.status_code,
            "body": external_response.json(),
        },
    }
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0 if response.status_code == 200 and external_response.status_code == 200 else 1


if __name__ == "__main__":
    raise SystemExit(main())
