from __future__ import annotations

import json
import os
import sys
import time
from datetime import datetime
from pathlib import Path
from uuid import uuid4

from fastapi.testclient import TestClient


def build_docx_bytes() -> bytes:
    from io import BytesIO

    from docx import Document as DocxDocument

    document = DocxDocument()
    document.add_heading("审批流程说明", level=1)
    document.add_paragraph("OA 文档说明：审批流程需要两级审批。")
    document.add_paragraph("第一步由部门负责人审批，第二步由 HR 系统负责人审批。")
    document.add_paragraph("审批通过后，流程状态会更新为已完成。")
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def build_markdown_report(result: dict[str, object]) -> str:
    checks = result["checks"]
    summary = checks.get("summary", {})
    search_hits = checks.get("search", {}).get("body", {}).get("data", {}).get("hits", [])
    top_hit = search_hits[0] if search_hits else {}
    qa_data = checks.get("qa", {}).get("body", {}).get("data", {}) or {}
    qa_answer = qa_data.get("answer", "N/A")
    qa_status = qa_data.get("answer_status", "not_available")
    return f"""# Knowledge Base Core Real Provider Smoke Report

**Date**: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
**Smoke ID**: {result["smoke_id"]}

## Providers

- Vector store: {result["providers"]["vector_store_provider"]}
- Embedding: {result["providers"]["embedding_provider"]}
- LLM: {result["providers"]["llm_provider"]}
- Zilliz collection: {result["providers"]["zilliz_collection"]}

## Checks

| Check | Status | Notes |
|---|---|---|
| health | PASS | HTTP {checks["health"]["status_code"]} |
| configs | PASS | HTTP {checks["configs"]["status_code"]} |
| upload | PASS | HTTP {checks.get("upload", {}).get("status_code", "N/A")} |
| job | PASS | HTTP {checks.get("job", {}).get("status_code", "N/A")} |
| detail | PASS | HTTP {checks.get("detail", {}).get("status_code", "N/A")} |
| download | PASS | {checks.get("download", {}).get("content_length", "N/A")} bytes |
| search | PASS | {checks.get("search", {}).get("hit_count", "N/A")} hits |
| debug_search | PASS | {checks.get("debug_search", {}).get("ranking_debug_count", "N/A")} ranking items |
| qa | PASS | {checks.get("qa", {}).get("citation_count", "N/A")} citations |
| cleanup | PASS | HTTP {checks.get("cleanup", {}).get("status_code", "N/A")} |

## Summary

- doc_uuid: `{summary.get("doc_uuid", "N/A")}`
- job_uuid: `{summary.get("job_uuid", "N/A")}`
- search_hits: {summary.get("search_hits", "N/A")}
- qa_citations: {summary.get("qa_citations", "N/A")}
- download_bytes: {summary.get("download_bytes", "N/A")}

## Retrieval Snapshot

- top title: {top_hit.get("title", "N/A")}
- top score: {top_hit.get("score", "N/A")}
- vector_score: {top_hit.get("vector_score", "N/A")}
- text_score: {top_hit.get("text_score", "N/A")}

## QA Snapshot

answer_status: `{qa_status}`

{qa_answer}

## Artifacts

- report_path: `{result["report_path"]}`
- artifact_path: `{result["artifact_path"]}`
"""


def build_output_paths(project_root: Path) -> tuple[Path, Path]:
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    base_dir = project_root.parent / "qa-outputs" / "oa-rag" / timestamp
    reports_dir = base_dir / "reports"
    artifacts_dir = base_dir / "artifacts"
    reports_dir.mkdir(parents=True, exist_ok=True)
    artifacts_dir.mkdir(parents=True, exist_ok=True)
    return reports_dir, artifacts_dir


def persist_result(*, project_root: Path, result: dict[str, object]) -> None:
    reports_dir, artifacts_dir = build_output_paths(project_root)
    report_path = reports_dir / "real-provider-smoke-report.md"
    artifact_path = artifacts_dir / "real-provider-smoke-result.json"
    result["report_path"] = str(report_path)
    result["artifact_path"] = str(artifact_path)
    report_path.write_text(build_markdown_report(result), encoding="utf-8")
    artifact_path.write_text(json.dumps(result, ensure_ascii=False, indent=2), encoding="utf-8")


def sync_smoke_identity(client: TestClient) -> tuple[bool, dict[str, object]]:
    source_response = client.post(
        "/api/v1/admin/identity-sources",
        json={
            "source_id": "real-provider-smoke-directory",
            "source_type": "json",
            "name": "Real Provider Smoke Directory",
            "field_mapping_json": {
                "user_id": "employee_no",
                "external_id": "directory_id",
                "display_name": "full_name",
                "email": "mail",
                "department_codes": "dept_codes",
                "status": "state",
            },
        },
    )
    if source_response.status_code not in {200, 409}:
        return False, {
            "source_status_code": source_response.status_code,
            "source_body": source_response.json(),
        }

    sync_response = client.post(
        "/api/v1/admin/identity-sources/real-provider-smoke-directory/sync",
        json={
            "records": [
                {
                    "employee_no": "smoke-user",
                    "directory_id": "real-provider-smoke-user",
                    "full_name": "Real Provider Smoke User",
                    "mail": "smoke-user@example.com",
                    "dept_codes": ["hr"],
                    "state": "active",
                }
            ]
        },
    )
    sync_body = sync_response.json()
    return sync_response.status_code == 200 and sync_body.get("data", {}).get("success_count") == 1, {
        "source_status_code": source_response.status_code,
        "source_body": source_response.json(),
        "sync_status_code": sync_response.status_code,
        "sync_body": sync_body,
    }


def main() -> int:
    project_root = Path(__file__).resolve().parents[1]
    if str(project_root) not in sys.path:
        sys.path.insert(0, str(project_root))

    smoke_id = uuid4().hex[:10]
    test_db_path = Path(f"/private/tmp/oa_rag_real_smoke_{smoke_id}.sqlite3")
    test_storage_root = Path(f"/private/tmp/oa_rag_real_smoke_data_{smoke_id}")
    smoke_collection = (
        os.environ.get("OA_RAG_SMOKE_COLLECTION")
        or os.environ.get("ZILLIZ_COLLECTION_NAME")
        or "oa_rag_chunks"
    ).strip() or "oa_rag_chunks"

    os.environ["DATABASE_URL"] = f"sqlite:///{test_db_path}"
    os.environ["STORAGE_ROOT"] = str(test_storage_root)
    os.environ["RAW_DATA_DIR"] = str(test_storage_root / "raw")
    os.environ["PROCESSED_DATA_DIR"] = str(test_storage_root / "processed")
    os.environ["SAMPLE_DATA_DIR"] = str(test_storage_root / "samples")
    os.environ["APP_API_KEY"] = "smoke-api-key"
    os.environ["INTERNAL_TOKEN"] = "smoke-internal-token"
    os.environ["ZILLIZ_COLLECTION_NAME"] = smoke_collection
    os.environ["HEALTH_PROBE_EXTERNAL_SERVICES"] = "true"
    os.environ["PROVIDER_RETRY_COUNT"] = os.environ.get("OA_RAG_SMOKE_PROVIDER_RETRY_COUNT", "3")

    from app.core.config import get_settings, reset_settings_cache
    from app.db.base import Base
    from app.db.runtime import get_engine, reset_db_runtime
    from app.main import create_app

    reset_settings_cache()
    reset_db_runtime()
    settings = get_settings()
    settings.ensure_storage_dirs()
    Base.metadata.create_all(bind=get_engine())

    app = create_app()
    client = TestClient(app)

    result: dict[str, object] = {
        "smoke_id": smoke_id,
        "providers": {
            "vector_store_provider": settings.vector_store_provider,
            "embedding_provider": settings.embedding_provider,
            "llm_provider": "http" if settings.llm_api_base and settings.llm_api_key and settings.llm_model else "local",
            "zilliz_collection": settings.zilliz_collection,
        },
        "checks": {},
    }
    doc_uuid: str | None = None

    def finalize_and_exit(exit_code: int) -> int:
        if doc_uuid:
            cleanup_resp = client.delete(f"/api/v1/documents/{doc_uuid}")
            result["checks"]["cleanup"] = {
                "status_code": cleanup_resp.status_code,
                "body": cleanup_resp.json(),
            }
        else:
            result["checks"]["cleanup"] = {
                "status_code": 0,
                "body": {"code": 0, "message": "skipped", "data": {"reason": "document_not_created"}},
            }
        persist_result(project_root=project_root, result=result)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return exit_code

    health_resp = client.get("/api/v1/health")
    result["checks"]["health"] = {
        "status_code": health_resp.status_code,
        "body": health_resp.json(),
    }
    if health_resp.status_code != 200:
        return finalize_and_exit(1)

    configs_resp = client.get("/api/v1/configs")
    result["checks"]["configs"] = {
        "status_code": configs_resp.status_code,
        "body": configs_resp.json(),
    }
    if configs_resp.status_code != 200:
        return finalize_and_exit(1)

    identity_ok, identity_detail = sync_smoke_identity(client)
    result["checks"]["identity_sync"] = identity_detail
    if not identity_ok:
        return finalize_and_exit(1)

    docx_bytes = build_docx_bytes()
    upload_resp = client.post(
        "/api/v1/documents/upload",
        data={
            "title": f"OA Smoke {smoke_id}",
            "source_type": "rule_doc",
            "source_module": "oa",
            "version": smoke_id,
            "access_level": "internal",
            "owner_dept": "hr",
            "tags": json.dumps(["smoke", "real-provider"]),
            "extra_meta": json.dumps({"scenario": "real-provider-smoke"}),
        },
        files={
            "file": (
                f"oa-smoke-{smoke_id}.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
    )
    upload_body = upload_resp.json()
    result["checks"]["upload"] = {
        "status_code": upload_resp.status_code,
        "body": upload_body,
    }
    if upload_resp.status_code != 200:
        return finalize_and_exit(1)

    job_uuid = upload_body["data"]["job_uuid"]
    doc_uuid = upload_body["data"]["doc_uuid"]

    job_resp = client.get(f"/api/v1/jobs/{job_uuid}")
    result["checks"]["job"] = {
        "status_code": job_resp.status_code,
        "body": job_resp.json(),
    }
    if job_resp.status_code != 200:
        return finalize_and_exit(1)

    detail_resp = client.get(f"/api/v1/documents/{doc_uuid}")
    detail_body = detail_resp.json()
    result["checks"]["detail"] = {
        "status_code": detail_resp.status_code,
        "body": detail_body,
    }
    if detail_resp.status_code != 200:
        return finalize_and_exit(1)

    download_resp = client.get(f"/api/v1/documents/{doc_uuid}/download")
    result["checks"]["download"] = {
        "status_code": download_resp.status_code,
        "content_length": len(download_resp.content),
    }
    if download_resp.status_code != 200 or not download_resp.content:
        return finalize_and_exit(1)

    search_resp = client.post(
        "/api/v1/retrieval/search",
        json={
            "query": "审批流程 两级审批",
            "top_k": 3,
            "filters": {"source_module": ["oa"]},
            "user_context": {"user_id": "smoke-user"},
        },
    )
    search_body = search_resp.json()
    result["checks"]["search"] = {
        "status_code": search_resp.status_code,
        "hit_count": len(search_body.get("data", {}).get("hits", [])),
        "body": search_body,
    }
    if search_resp.status_code != 200 or not search_body.get("data", {}).get("hits"):
        return finalize_and_exit(1)

    debug_search_resp = client.post(
        "/api/v1/retrieval/debug-search",
        json={
            "query": "审批流程 两级审批",
            "top_k": 3,
            "filters": {"source_module": ["oa"]},
            "user_context": {"user_id": "smoke-user"},
        },
    )
    debug_search_body = debug_search_resp.json()
    result["checks"]["debug_search"] = {
        "status_code": debug_search_resp.status_code,
        "ranking_debug_count": len(debug_search_body.get("data", {}).get("ranking_debug", [])),
        "body": debug_search_body,
    }
    if debug_search_resp.status_code != 200:
        return finalize_and_exit(1)

    time.sleep(0.5)

    qa_resp = client.post(
        "/api/v1/qa/answer",
        headers={"X-Internal-Token": "smoke-internal-token"},
        json={
            "question": "审批流程 两级审批 部门负责人 HR 系统负责人 是什么？",
            "top_k": 3,
            "filters": {"source_module": ["oa"]},
            "user_context": {"user_id": "smoke-user"},
            "generation_options": {"temperature": 0.2, "max_tokens": 200},
        },
    )
    qa_body = qa_resp.json()
    qa_data = qa_body.get("data") if isinstance(qa_body.get("data"), dict) else {}
    result["checks"]["qa"] = {
        "status_code": qa_resp.status_code,
        "citation_count": len(qa_data.get("citations", [])),
        "answer_preview": str(qa_data.get("answer", ""))[:160],
        "body": qa_body,
    }
    if (
        qa_resp.status_code != 200
        or not qa_data.get("answer")
        or qa_data.get("answer_status") != "grounded"
    ):
        return finalize_and_exit(1)

    result["checks"]["summary"] = {
        "doc_uuid": doc_uuid,
        "job_uuid": job_uuid,
        "search_hits": len(search_body["data"]["hits"]),
        "qa_citations": len(qa_body["data"]["citations"]),
        "download_bytes": len(download_resp.content),
    }

    cleanup_resp = client.delete(f"/api/v1/documents/{doc_uuid}")
    result["checks"]["cleanup"] = {
        "status_code": cleanup_resp.status_code,
        "body": cleanup_resp.json(),
    }
    doc_uuid = None
    if cleanup_resp.status_code != 200:
        persist_result(project_root=project_root, result=result)
        print(json.dumps(result, ensure_ascii=False, indent=2))
        return 1

    persist_result(project_root=project_root, result=result)
    print(json.dumps(result, ensure_ascii=False, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
